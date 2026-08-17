from datetime import datetime, date
from io import BytesIO
import os
import sqlite3
import pandas as pd
import streamlit as st
from PIL import Image, ImageOps
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
)
from docx import Document

# =========================================================
# CONFIGURATION & THÈME BLEU / BLANC / VERT / ROUGE
# =========================================================
st.set_page_config(
    page_title="Gestion Stock MW NOMATIS",
    page_icon="📡",
    layout="wide",
    initial_sidebar_state="collapsed",
)

APP_TITLE = "Gestion Stock MW NOMATIS"
DB_FILE = "stock_mw.db"

CLIENTS = {
    "Orange": {"logo": "Orange_logo.svg.webp", "color": "#FF6600"},
    "Inwi": {"logo": "Logo INWI.jpg", "color": "#A1006B"},
    "ZTE": {"logo": "Logo ZTE.jpg", "color": "#005BAC"},
}

DEFAULT_ARTICLES = ["Câble IF", "Câble RJ45", "Support 0.3 m", "Support 0.6 m", "ODU 18GHz", "Antenne 0.6m"]
DEFAULT_FOURNISSEURS = ["NEC", "ZTE", "Intégral", "FO Connect"]
DEFAULT_EQUIPES = ["Nabil Team", "Yassine Team", "Issam Team"]

PERMISSIONS = {
    "admin": {"be", "bs", "stock", "edit", "config", "manage_users"},
    "magasinier": {"be", "bs", "stock", "edit"},
    "coordinateur": {"stock", "print"},
    "coordinatrice": {"stock", "print"},
}

# Injection CSS — Respect strict des couleurs
st.markdown(
    """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

        html, body, [class*="css"] {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
            background-color: #F4F7FA !important;
            color: #0F172A !important;
        }

        .main-header {
            background: linear-gradient(135deg, #1E3A8A 0%, #2563EB 100%);
            border-bottom: 4px solid #10B981;
            border-radius: 12px;
            padding: 20px 28px;
            margin-bottom: 24px;
            color: #FFFFFF !important;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        }

        .main-title {
            color: #FFFFFF !important;
            font-weight: 700;
            font-size: 26px;
            margin: 0;
        }

        .subtitle {
            color: #E2E8F0 !important;
            font-size: 14px;
            margin-top: 4px;
        }

        .glass-card {
            background: #FFFFFF;
            border: 1px solid #E2E8F0;
            border-radius: 12px;
            padding: 24px;
            margin-bottom: 20px;
            box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.05);
        }

        /* Styles spécifiques aux boutons */
        .btn-login > button, div[data-testid="stFormSubmitButton"] > button {
            background-color: #DC2626 !important; /* Rouge au départ */
            color: #FFFFFF !important;
            border-radius: 8px !important;
            font-weight: 700 !important;
            border: none !important;
            min-height: 45px !important;
            width: 100%;
        }

        .btn-login-success > button {
            background-color: #10B981 !important; /* Vert après connexion */
            color: #FFFFFF !important;
            border-radius: 8px !important;
            font-weight: 700 !important;
            border: none !important;
            min-height: 45px !important;
            width: 100%;
        }

        .stButton > button {
            background-color: #2563EB;
            color: #FFFFFF;
            border-radius: 8px;
            font-weight: 600;
            border: none;
            padding: 8px 16px;
        }

        div[data-baseweb="tab-list"] {
            gap: 8px;
            background-color: #E2E8F0;
            padding: 6px;
            border-radius: 10px;
        }

        button[data-baseweb="tab"] {
            border-radius: 6px !important;
            font-weight: 600 !important;
            color: #1E3A8A !important;
        }

        button[aria-selected="true"] {
            background-color: #10B981 !important; /* Vert Accent */
            color: #FFFFFF !important;
        }
    </style>
    """,
    unsafe_allow_html=True,
)


# =========================================================
# BASE DE DONNÉES (SQLite)
# =========================================================
def get_conn():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_conn()
    cur = conn.cursor()
    cur.executescript(
        """
        CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            password TEXT NOT NULL,
            fullname TEXT NOT NULL,
            role TEXT NOT NULL,
            last_login TEXT
        );

        CREATE TABLE IF NOT EXISTS articles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT UNIQUE NOT NULL,
            active INTEGER NOT NULL DEFAULT 1
        );

        CREATE TABLE IF NOT EXISTS fournisseurs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT UNIQUE NOT NULL,
            active INTEGER NOT NULL DEFAULT 1
        );

        CREATE TABLE IF NOT EXISTS equipes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT UNIQUE NOT NULL,
            active INTEGER NOT NULL DEFAULT 1
        );

        CREATE TABLE IF NOT EXISTS stock (
            client TEXT NOT NULL,
            article_id INTEGER NOT NULL,
            quantity INTEGER NOT NULL DEFAULT 0,
            PRIMARY KEY (client, article_id)
        );

        CREATE TABLE IF NOT EXISTS bons (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            type TEXT NOT NULL,
            number TEXT NOT NULL,
            client TEXT NOT NULL,
            date_bon TEXT NOT NULL,
            datetime_saisie TEXT NOT NULL,
            fournisseur TEXT,
            lieu_livraison TEXT,
            receptionne_par TEXT,
            equipe TEXT,
            destination TEXT,
            created_by TEXT NOT NULL,
            UNIQUE(type, number, client)
        );

        CREATE TABLE IF NOT EXISTS bon_items (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            bon_id INTEGER NOT NULL,
            article_id INTEGER NOT NULL,
            reference TEXT,
            quantity INTEGER NOT NULL,
            remarque TEXT,
            FOREIGN KEY(bon_id) REFERENCES bons(id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS movements (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            client TEXT NOT NULL,
            article_id INTEGER NOT NULL,
            movement_type TEXT NOT NULL,
            quantity INTEGER NOT NULL,
            reference_bon TEXT,
            username TEXT NOT NULL,
            created_at TEXT NOT NULL,
            comment TEXT,
            fournisseur TEXT,
            equipe TEXT
        );
        """
    )

    cur.execute("SELECT COUNT(*) FROM users")
    if cur.fetchone()[0] == 0:
        cur.execute(
            "INSERT INTO users VALUES (?,?,?,?,?)",
            ("admin", "admin123", "Administrateur Système", "admin", "Jamais"),
        )
        cur.execute(
            "INSERT INTO users VALUES (?,?,?,?,?)",
            ("magasinier", "123", "Magasinier Principal", "magasinier", "Jamais"),
        )
        cur.execute(
            "INSERT INTO users VALUES (?,?,?,?,?)",
            ("coord", "123", "Coordinateur Projet", "coordinateur", "Jamais"),
        )

    for name in DEFAULT_ARTICLES:
        cur.execute("INSERT OR IGNORE INTO articles(name,active) VALUES(?,1)", (name,))
    for name in DEFAULT_FOURNISSEURS:
        cur.execute("INSERT OR IGNORE INTO fournisseurs(name,active) VALUES(?,1)", (name,))
    for name in DEFAULT_EQUIPES:
        cur.execute("INSERT OR IGNORE INTO equipes(name,active) VALUES(?,1)", (name,))

    conn.commit()
    conn.close()


init_db()


# =========================================================
# UTILITAIRES & REQUÊTES
# =========================================================
def query(sql, params=(), one=False):
    conn = get_conn()
    cur = conn.execute(sql, params)
    rows = cur.fetchall()
    conn.close()
    return rows[0] if one and rows else (None if one else rows)


def execute(sql, params=()):
    conn = get_conn()
    cur = conn.execute(sql, params)
    conn.commit()
    last_id = cur.lastrowid
    conn.close()
    return last_id


def active_names(table):
    rows = query(f"SELECT name FROM {table} WHERE active=1 ORDER BY name")
    return [r["name"] for r in rows]


def article_id_by_name(name):
    row = query("SELECT id FROM articles WHERE name=?", (name,), one=True)
    return row["id"] if row else None


def current_stock(client, article_id):
    row = query("SELECT quantity FROM stock WHERE client=? AND article_id=?", (client, article_id), one=True)
    return int(row["quantity"]) if row else 0


def set_stock(client, article_id, quantity):
    execute(
        """
        INSERT INTO stock(client,article_id,quantity) VALUES(?,?,?)
        ON CONFLICT(client,article_id) DO UPDATE SET quantity=excluded.quantity
        """,
        (client, article_id, max(0, int(quantity))),
    )


def add_movement(client, article_id, m_type, qty, ref_bon, user, comment="", fournisseur="", equipe=""):
    execute(
        """
        INSERT INTO movements(client,article_id,movement_type,quantity,reference_bon,username,created_at,comment,fournisseur,equipe)
        VALUES(?,?,?,?,?,?,?,?,?,?)
        """,
        (client, article_id, m_type, int(qty), ref_bon, user, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), comment, fournisseur, equipe),
    )


def user_info(username):
    return query("SELECT * FROM users WHERE username=?", (username,), one=True)


def can(role, permission):
    return permission in PERMISSIONS.get(role, set())


def normalized_logo(path, size=(250, 100)):
    if not os.path.exists(path):
        return None
    try:
        img = Image.open(path).convert("RGB")
        canvas = Image.new("RGB", size, "#FFFFFF")
        contained = ImageOps.contain(img, size)
        x = (size[0] - contained.width) // 2
        y = (size[1] - contained.height) // 2
        canvas.paste(contained, (x, y))
        return canvas
    except Exception:
        return None


# =========================================================
# MODALES DE CONFIRMATION (DIALOG)
# =========================================================
@st.dialog("⚠️ Confirmation requise")
def confirm_action_dialog(message, callback, *args):
    st.write(message)
    col1, col2 = st.columns(2)
    if col1.button("✅ Oui, Confirmer", use_container_width=True):
        callback(*args)
        st.rerun()
    if col2.button("❌ Annuler", use_container_width=True):
        st.rerun()


# =========================================================
# GÉNÉRATION DE DOCUMENTS (PDF & DOCX)
# =========================================================
def generate_pdf(bon_id):
    bon = query("SELECT * FROM bons WHERE id=?", (bon_id,), one=True)
    items = query(
        "SELECT bi.*, a.name AS article FROM bon_items bi JOIN articles a ON a.id=bi.article_id WHERE bi.bon_id=?",
        (bon_id,),
    )
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=15 * mm, leftMargin=15 * mm, topMargin=15 * mm, bottomMargin=15 * mm)
    styles = getSampleStyleSheet()
    story = []

    logo_info = CLIENTS.get(bon["client"], {})
    logo_path = logo_info.get("logo")

    header_data = []
    if logo_path and os.path.exists(logo_path):
        try:
            img_rl = RLImage(logo_path, width=40 * mm, height=18 * mm)
            header_data.append([img_rl, Paragraph("<font size=16 color='#1E3A8A'><b>Gestion Stock MW NOMATIS</b></font>", styles["Normal"])])
        except Exception:
            header_data.append([Paragraph("<font size=16 color='#1E3A8A'><b>NOMATIS</b></font>", styles["Normal"])])
    else:
        header_data.append([Paragraph("<font size=16 color='#1E3A8A'><b>NOMATIS</b></font>", styles["Normal"])])

    header_table = Table(header_data, colWidths=[50 * mm, 130 * mm])
    header_table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    story.append(header_table)
    story.append(Spacer(1, 10))

    title = "BON D'ENTRÉE (BE)" if bon["type"] == "BE" else "BON DE SORTIE (BS)"
    story.append(Paragraph(f"<font size=18 color='#2563EB'><b>{title} — {bon['number']}</b></font>", styles["Title"]))
    story.append(Spacer(1, 10))

    if bon["type"] == "BE":
        info = [
            ["N° Bon", bon["number"], "Date / Heure Saisie", bon["datetime_saisie"]],
            ["Date du Bon", bon["date_bon"], "Fournisseur", bon["fournisseur"] or ""],
            ["Réceptionné par", bon["receptionne_par"] or "", "Client / Projet", bon["client"]],
            ["Lieu Livraison", bon["lieu_livraison"] or "", "", ""],
        ]
    else:
        info = [
            ["N° Bon", bon["number"], "Date / Heure Saisie", bon["datetime_saisie"]],
            ["Date du Bon", bon["date_bon"], "Équipe Destination", bon["equipe"] or ""],
            ["Saisi par", bon["created_by"], "Destination / Site", bon["destination"] or ""],
            ["Client / Projet", bon["client"], "", ""],
        ]

    info_table = Table(info, colWidths=[35 * mm, 55 * mm, 35 * mm, 55 * mm])
    info_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F1F5F9")),
                ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#F1F5F9")),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
                ("PADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(info_table)
    story.append(Spacer(1, 15))

    data = [["Référence", "Désignation Article", "Quantité", "Remarque"]]
    for item in items:
        data.append([item["reference"] or "-", item["article"], str(item["quantity"]), item["remarque"] or "-"])

    items_table = Table(data, colWidths=[35 * mm, 70 * mm, 25 * mm, 50 * mm])
    items_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563EB")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
                ("ALIGN", (2, 1), (2, -1), "CENTER"),
                ("PADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(items_table)
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_docx(bon_id):
    bon = query("SELECT * FROM bons WHERE id=?", (bon_id,), one=True)
    items = query(
        "SELECT bi.*, a.name AS article FROM bon_items bi JOIN articles a ON a.id=bi.article_id WHERE bi.bon_id=?",
        (bon_id,),
    )
    doc = Document()
    doc.add_heading(f"BON DE {'ENTRÉE' if bon['type']=='BE' else 'SORTIE'} - {bon['number']}", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Client / Projet : {bon['client']}\nDate Bon : {bon['date_bon']}\nSaisie le : {bon['datetime_saisie']}\nOpérateur : {bon['created_by']}")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Référence"
    hdr_cells[1].text = "Article"
    hdr_cells[2].text = "Quantité"
    hdr_cells[3].text = "Remarque"

    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item["reference"] or ""
        row_cells[1].text = item["article"]
        row_cells[2].text = str(item["quantity"])
        row_cells[3].text = item["remarque"] or ""

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# =========================================================
# SESSION & AUTHENTIFICATION
# =========================================================
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "current_user" not in st.session_state:
    st.session_state.current_user = None
if "selected_client" not in st.session_state:
    st.session_state.selected_client = None
if "temp_be_items" not in st.session_state:
    st.session_state.temp_be_items = []
if "temp_bs_items" not in st.session_state:
    st.session_state.temp_bs_items = []


def login_screen():
    st.markdown("<br>", unsafe_allow_html=True)
    _, col, _ = st.columns([1, 1.8, 1])
    with col:
        st.markdown(
            """
            <div class="glass-card" style="text-align: center; border-top: 5px solid #2563EB;">
                <h1 style="color: #1E3A8A; margin-bottom: 0px; font-weight: 800;">Gestion Stock MW NOMATIS</h1>
                <p style="color: #10B981; font-weight: 600; font-size: 15px;">Plateforme Sécurisée d'Ingénierie & Logistique</p>
            </div>
            """,
            unsafe_allow_html=True,
        )
        with st.form("login_form"):
            st.markdown("##### Connexion Utilisateur")
            username = st.text_input("Nom d'utilisateur")
            password = st.text_input("Mot de passe", type="password")

            btn_class = "btn-login-success" if st.session_state.logged_in else "btn-login"
            st.markdown(f'<div class="{btn_class}">', unsafe_allow_html=True)
            submit = st.form_submit_button("SE CONNECTER", use_container_width=True)
            st.markdown("</div>", unsafe_allow_html=True)

            if submit:
                user = user_info(username)
                if user and user["password"] == password:
                    execute("UPDATE users SET last_login=? WHERE username=?", (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), username))
                    st.session_state.logged_in = True
                    st.session_state.current_user = username
                    st.success("Accès validé !")
                    st.rerun()
                else:
                    st.error("Identifiants incorrects.")


if not st.session_state.logged_in:
    login_screen()
    st.stop()

CURRENT_USER = user_info(st.session_state.current_user)
ROLE = CURRENT_USER["role"]


# =========================================================
# SELECTION CLIENT & GESTION DU COMPTE USER / ADMIN
# =========================================================
if not st.session_state.selected_client:
    st.markdown(
        """
        <div class="main-header">
            <div class="main-title">Gestion Stock MW NOMATIS</div>
            <div class="subtitle">Sélectionnez l'espace client pour accéder au stock dédié</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Grille des clients
    cols = st.columns(3)
    for idx, (client, info) in enumerate(CLIENTS.items()):
        with cols[idx]:
            st.markdown(f'<div class="glass-card" style="border-top: 4px solid {info["color"]}; text-align: center;">', unsafe_allow_html=True)
            logo = normalized_logo(info["logo"])
            if logo:
                st.image(logo, use_container_width=True)
            else:
                st.markdown(f"<h2 style='color:{info['color']}'>{client}</h2>", unsafe_allow_html=True)

            # Bouton en gras avec la couleur du logo
            st.markdown(
                f"""
                <style>
                    div[data-testid="stBlock"] button[key="select_{client}"] {{
                        background-color: {info['color']} !important;
                        font-weight: bold !important;
                        color: #FFFFFF !important;
                    }}
                </style>
                """,
                unsafe_allow_html=True,
            )
            if st.button(f"ACCÈS AU STOCK {client.upper()}", key=f"select_{client}", use_container_width=True):
                st.session_state.selected_client = client
                st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("---")

    # Section Profil / Administration des Utilisateurs
    if can(ROLE, "manage_users"):
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.subheader("⚙️ Administration des Utilisateurs (Rôle Admin)")

        tab_u1, tab_u2 = st.tabs(["➕ Créer un Utilisateur", "✏️ Gérer / Modifier Utilisateurs"])

        with tab_u1:
            with st.form("create_user_form"):
                new_u = st.text_input("Identifiant (Username)")
                new_p = st.text_input("Mot de passe", type="password")
                new_fn = st.text_input("Nom Complet")
                new_r = st.selectbox("Rôle", ["admin", "magasinier", "coordinateur", "coordinatrice"])
                if st.form_submit_button("Créer l'utilisateur"):
                    if new_u and new_p and new_fn:
                        try:
                            execute("INSERT INTO users VALUES (?,?,?,?,?)", (new_u, new_p, new_fn, new_r, "Jamais"))
                            st.success(f"Utilisateur {new_u} créé !")
                            st.rerun()
                        except Exception:
                            st.error("Nom d'utilisateur déjà existant.")
                    else:
                        st.error("Veuillez remplir tous les champs.")

        with tab_u2:
            all_users = query("SELECT username, fullname, role, last_login FROM users")
            st.dataframe(pd.DataFrame([dict(u) for u in all_users]), use_container_width=True, hide_index=True)

            u_to_edit = st.selectbox("Sélectionner un utilisateur à modifier", [u["username"] for u in all_users])
            if u_to_edit:
                u_data = user_info(u_to_edit)
                with st.form("edit_user_form"):
                    e_fn = st.text_input("Nom complet", value=u_data["fullname"])
                    e_p = st.text_input("Nouveau mot de passe", value=u_data["password"])
                    e_r = st.selectbox("Rôle", ["admin", "magasinier", "coordinateur", "coordinatrice"], index=["admin", "magasinier", "coordinateur", "coordinatrice"].index(u_data["role"]))
                    if st.form_submit_button("Mettre à jour l'utilisateur"):
                        execute("UPDATE users SET fullname=?, password=?, role=? WHERE username=?", (e_fn, e_p, e_r, u_to_edit))
                        st.success("Informations mises à jour !")
                        st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    else:
        # Utilisateur normal : Modification de son propre compte
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.subheader("👤 Mon Compte")
        with st.form("my_account_form"):
            my_fn = st.text_input("Nom Complet", value=CURRENT_USER["fullname"])
            my_p = st.text_input("Changer le mot de passe", value=CURRENT_USER["password"], type="password")
            if st.form_submit_button("Mettre à jour mes informations"):
                execute("UPDATE users SET fullname=?, password=? WHERE username=?", (my_fn, my_p, CURRENT_USER["username"]))
                st.success("Compte mis à jour avec succès !")
                st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    st.stop()

CLIENT = st.session_state.selected_client


# =========================================================
# APPLI PRINCIPALE : 5 RUBRIQUES
# =========================================================
h1, h2 = st.columns([3, 1])
with h1:
    st.markdown(
        f"""
        <div class="main-header">
            <div class="main-title">Gestion Stock MW NOMATIS — Espaces {CLIENT}</div>
            <div class="subtitle">Utilisateur : {CURRENT_USER['fullname']} ({ROLE.upper()})</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
with h2:
    if st.button("🔄 Changer Client", use_container_width=True):
        st.session_state.selected_client = None
        st.rerun()
    if st.button("🚪 Déconnexion", use_container_width=True):
        st.session_state.logged_in = False
        st.session_state.selected_client = None
        st.rerun()

# 5 Rubriques obligatoires
tabs = st.tabs(["📥 BE", "📤 BS", "📊 Situation Stock", "📜 Historique", "⚙️ Configuration"])

# ---------------------------------------------------------
# 📥 RUBRIQUE 1: BON D'ENTRÉE (BE)
# ---------------------------------------------------------
with tabs[0]:
    if can(ROLE, "be"):
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.subheader("Création Bon d'Entrée (BE)")

        c1, c2, c3 = st.columns(3)
        now_dt = datetime.now()
        c1.text_input("Date/Heure Saisie (Auto)", value=now_dt.strftime("%Y-%m-%d %H:%M:%S"), disabled=True)
        date_be = c2.date_input("Date Bon (<= Aujourd'hui)", value=date.today(), max_value=date.today())

        auto_num_be = f"BE MW-{now_dt.strftime('%Y%m%d')}-01"
        num_be = c3.text_input("N° BE", value=auto_num_be)

        c4, c5 = st.columns(2)
        existing_fourns = active_names("fournisseurs")
        fourn_sel = c4.selectbox("Fournisseur existant", ["-- Autre --"] + existing_fourns)
        fourn_custom = c4.text_input("Saisir autre fournisseur (si non listé)")
        fournisseur = fourn_custom if fourn_sel == "-- Autre --" else fourn_sel

        lieu = c5.text_input("Lieu de Livraison", value="Magasin Principal")
        st.text_input("Réceptionné par (Auto)", value=CURRENT_USER["fullname"], disabled=True)

        st.markdown("---")
        st.markdown("##### Ajouter des articles au Bon")
        r1, r2, r3, r4 = st.columns([2, 3, 1.5, 3])
        ref = r1.text_input("Référence", key="be_ref")
        art = r2.selectbox("Article (Prédéfini)", active_names("articles"), key="be_art")
        qty = r3.number_input("Quantité (> 0)", min_value=1, value=1, key="be_qty")
        rem = r4.text_input("Remarque", key="be_rem")

        if st.button("➕ Ajouter au tableau", key="add_be_line"):
            if not fournisseur:
                st.error("Le champ Fournisseur est obligatoire.")
            else:
                # Regroupement si même article
                found = False
                for item in st.session_state.temp_be_items:
                    if item["art"] == art and item["ref"] == ref:
                        item["qty"] += qty
                        item["rem"] = (item["rem"] + " | " + rem).strip(" | ")
                        found = True
                        break
                if not found:
                    st.session_state.temp_be_items.append({"ref": ref, "art": art, "qty": qty, "rem": rem})

        if st.session_state.temp_be_items:
            st.markdown("###### Articles saisis :")
            st.dataframe(pd.DataFrame(st.session_state.temp_be_items), use_container_width=True)

            def save_be():
                bon_id = execute(
                    "INSERT INTO bons (type,number,client,date_bon,datetime_saisie,fournisseur,lieu_livraison,receptionne_par,created_by) VALUES (?,?,?,?,?,?,?,?,?)",
                    (
                        "BE",
                        num_be,
                        CLIENT,
                        str(date_be),
                        now_dt.strftime("%Y-%m-%d %H:%M:%S"),
                        fournisseur,
                        lieu,
                        CURRENT_USER["fullname"],
                        CURRENT_USER["username"],
                    ),
                )
                for item in st.session_state.temp_be_items:
                    art_id = article_id_by_name(item["art"])
                    execute("INSERT INTO bon_items (bon_id,article_id,reference,quantity,remarque) VALUES (?,?,?,?,?)", (bon_id, art_id, item["ref"], item["qty"], item["rem"]))
                    set_stock(CLIENT, art_id, current_stock(CLIENT, art_id) + item["qty"])
                    add_movement(CLIENT, art_id, "BE", item["qty"], num_be, CURRENT_USER["username"], item["rem"], fournisseur=fournisseur)
                st.session_state.temp_be_items = []
                st.success("Bon d'Entrée enregistré !")

            b1, b2 = st.columns(2)
            if b1.button("💾 Enregistrer le Bon", use_container_width=True):
                confirm_action_dialog("Voulez-vous vraiment enregistrer ce Bon d'Entrée ?", save_be)
            if b2.button("🗑️ Vider le tableau", use_container_width=True):
                st.session_state.temp_be_items = []
                st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)
    else:
        st.info("Accès restreint : Consultez la Situation Stock ou l'Historique.")

# ---------------------------------------------------------
# 📤 RUBRIQUE 2: BON DE SORTIE (BS)
# ---------------------------------------------------------
with tabs[1]:
    if can(ROLE, "bs"):
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.subheader("Création Bon de Sortie (BS)")

        c1, c2, c3 = st.columns(3)
        now_dt = datetime.now()
        c1.text_input("Date/Heure Saisie (Auto)", value=now_dt.strftime("%Y-%m-%d %H:%M:%S"), disabled=True, key="bs_dt")
        date_bs = c2.date_input("Date Bon (<= Aujourd'hui)", value=date.today(), max_value=date.today(), key="bs_d")

        auto_num_bs = f"BS MW-{now_dt.strftime('%Y%m%d')}-01"
        num_bs = c3.text_input("N° BS", value=auto_num_bs)

        c4, c5 = st.columns(2)
        equipe = c4.selectbox("Équipe Récupératrice", active_names("equipes"))
        destination = c5.text_input("Destination / Site", value="Site Telecom")

        st.markdown("---")
        st.markdown("##### Sélection des articles à sortir")
        r1, r2, r3, r4 = st.columns([2, 3, 1.5, 3])
        ref = r1.text_input("Référence", key="bs_ref")
        art = r2.selectbox("Article", active_names("articles"), key="bs_art")
        qty = r3.number_input("Quantité", min_value=1, value=1, key="bs_qty")
        rem = r4.text_input("Remarque", key="bs_rem")

        if st.button("➕ Ajouter au BS", key="add_bs_line"):
            art_id = article_id_by_name(art)
            stk_dispo = current_stock(CLIENT, art_id)

            # Calcul quantité déjà en panier temporaire
            in_cart = sum(item["qty"] for item in st.session_state.temp_bs_items if item["art"] == art)
            if (qty + in_cart) > stk_dispo:
                st.error(f"Quantité insuffisante ! Stock dispo : {stk_dispo} (En panier: {in_cart})")
            else:
                found = False
                for item in st.session_state.temp_bs_items:
                    if item["art"] == art and item["ref"] == ref:
                        item["qty"] += qty
                        item["rem"] = (item["rem"] + " | " + rem).strip(" | ")
                        found = True
                        break
                if not found:
                    st.session_state.temp_bs_items.append({"ref": ref, "art": art, "qty": qty, "rem": rem})

        if st.session_state.temp_bs_items:
            st.markdown("###### Articles dans le panier de sortie :")
            st.dataframe(pd.DataFrame(st.session_state.temp_bs_items), use_container_width=True)

            def save_bs():
                bon_id = execute(
                    "INSERT INTO bons (type,number,client,date_bon,datetime_saisie,equipe,destination,created_by) VALUES (?,?,?,?,?,?,?,?)",
                    (
                        "BS",
                        num_bs,
                        CLIENT,
                        str(date_bs),
                        now_dt.strftime("%Y-%m-%d %H:%M:%S"),
                        equipe,
                        destination,
                        CURRENT_USER["username"],
                    ),
                )
                for item in st.session_state.temp_bs_items:
                    art_id = article_id_by_name(item["art"])
                    execute("INSERT INTO bon_items (bon_id,article_id,reference,quantity,remarque) VALUES (?,?,?,?,?)", (bon_id, art_id, item["ref"], item["qty"], item["rem"]))
                    set_stock(CLIENT, art_id, current_stock(CLIENT, art_id) - item["qty"])
                    add_movement(CLIENT, art_id, "BS", item["qty"], num_bs, CURRENT_USER["username"], item["rem"], equipe=equipe)
                st.session_state.temp_bs_items = []
                st.success("Bon de Sortie enregistré !")

            b1, b2 = st.columns(2)
            if b1.button("💾 Enregistrer le BS", use_container_width=True):
                confirm_action_dialog("Voulez-vous vraiment enregistrer ce Bon de Sortie ?", save_bs)
            if b2.button("🗑️ Vider la liste", use_container_width=True):
                st.session_state.temp_bs_items = []
                st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)
    else:
        st.info("Accès restreint.")

# ---------------------------------------------------------
# 📊 RUBRIQUE 3: SITUATION STOCK
# ---------------------------------------------------------
with tabs[2]:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader(f"Situation du Stock en Temps Réel — {CLIENT}")

    rows = query(
        """
        SELECT 
            a.name AS Article,
            COALESCE((SELECT SUM(m.quantity) FROM movements m WHERE m.article_id = a.id AND m.client = ? AND m.movement_type IN ('BE','AJUST_POS')), 0) AS Total_Entrees,
            COALESCE((SELECT SUM(m.quantity) FROM movements m WHERE m.article_id = a.id AND m.client = ? AND m.movement_type IN ('BS','AJUST_NEG')), 0) AS Total_Sorties,
            COALESCE(s.quantity, 0) AS Stock_Actuel
        FROM articles a
        LEFT JOIN stock s ON s.article_id = a.id AND s.client = ?
        WHERE a.active = 1 ORDER BY a.name
        """,
        (CLIENT, CLIENT, CLIENT),
    )
    df_stock = pd.DataFrame([dict(r) for r in rows])

    if not df_stock.empty:
        c1, c2 = st.columns(2)
        c1.metric("Nombre d'Articles Référencés", len(df_stock))
        c2.metric("Total Unités en Stock", int(df_stock["Stock_Actuel"].sum()))
        st.dataframe(df_stock, use_container_width=True, hide_index=True)

        # Impression / Export
        st.download_button(
            "🖨️ Exporter / Imprimer Situation Stock (CSV)",
            df_stock.to_csv(index=False).encode("utf-8"),
            file_name=f"Situation_Stock_{CLIENT}_{date.today()}.csv",
            mime="text/csv",
        )
    st.markdown("</div>", unsafe_allow_html=True)

# ---------------------------------------------------------
# 📜 RUBRIQUE 4: HISTORIQUE & MODIFICATION
# ---------------------------------------------------------
with tabs[3]:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("Historique des Bons & Impression")

    type_filter = st.radio("Type de Bon :", ["Bons d'Entrée (BE)", "Bons de Sortie (BS)"], horizontal=True)
    b_type = "BE" if "Entrée" in type_filter else "BS"

    bons = query("SELECT * FROM bons WHERE client=? AND type=? ORDER BY id DESC", (CLIENT, b_type))

    if bons:
        opts = [f"{b['number']} | Date: {b['date_bon']} | ID:{b['id']}" for b in bons]
        sel = st.selectbox("Sélectionner un Bon", opts)
        selected_id = int(sel.split("ID:")[1])
        b_data = query("SELECT * FROM bons WHERE id=?", (selected_id,), one=True)

        c1, c2 = st.columns(2)
        with c1:
            pdf_b = generate_pdf(selected_id)
            st.download_button("📄 Imprimer / Télécharger en PDF", pdf_b, file_name=f"{b_data['number']}.pdf", mime="application/pdf", use_container_width=True)
        with c2:
            docx_b = generate_docx(selected_id)
            st.download_button("📝 Imprimer / Télécharger en Word (.docx)", docx_b, file_name=f"{b_data['number']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

        # Modification & Suppression (Magasinier & Admin)
        if can(ROLE, "edit"):
            st.markdown("---")
            st.markdown("##### Options de Modification / Suppression")

            with st.expander("✏️ Modifier les informations du Bon"):
                with st.form("edit_bon_meta"):
                    n_date = st.date_input("Date Bon", value=datetime.strptime(b_data["date_bon"], "%Y-%m-%d").date(), max_value=date.today())
                    if b_type == "BE":
                        n_fourn = st.text_input("Fournisseur", value=b_data["fournisseur"] or "")
                        n_lieu = st.text_input("Lieu Livraison", value=b_data["lieu_livraison"] or "")
                        if st.form_submit_button("Valider modifications"):
                            execute("UPDATE bons SET date_bon=?, fournisseur=?, lieu_livraison=? WHERE id=?", (str(n_date), n_fourn, n_lieu, selected_id))
                            st.success("Modifié avec succès !")
                            st.rerun()
                    else:
                        n_eq = st.selectbox("Équipe", active_names("equipes"))
                        n_dest = st.text_input("Destination", value=b_data["destination"] or "")
                        if st.form_submit_button("Valider modifications"):
                            execute("UPDATE bons SET date_bon=?, equipe=?, destination=? WHERE id=?", (str(n_date), n_eq, n_dest, selected_id))
                            st.success("Modifié avec succès !")
                            st.rerun()

            def delete_current_bon():
                items_to_revert = query("SELECT * FROM bon_items WHERE bon_id=?", (selected_id,))
                for it in items_to_revert:
                    stk = current_stock(CLIENT, it["article_id"])
                    # Restauration stock
                    new_qty = stk - it["quantity"] if b_type == "BE" else stk + it["quantity"]
                    set_stock(CLIENT, it["article_id"], new_qty)
                execute("DELETE FROM bons WHERE id=?", (selected_id,))
                st.success("Bon supprimé et stock réajusté !")

            if st.button("🚨 Supprimer ce Bon (Restaure le stock)", use_container_width=True):
                confirm_action_dialog("Êtes-vous sûr de vouloir SUPPRIMER ce Bon ? Cette action réajustera le stock.", delete_current_bon)

    st.markdown("---")
    st.subheader("Historique Détaillé des Mouvements")

    f_fourn = st.selectbox("Filtrer par Fournisseur", ["Tous"] + active_names("fournisseurs"))
    f_eq = st.selectbox("Filtrer par Équipe", ["Tous"] + active_names("equipes"))

    sql_m = "SELECT m.created_at AS Date, m.movement_type AS Type, m.reference_bon AS Bon, a.name AS Article, m.quantity AS Qte, m.fournisseur, m.equipe, m.username AS Operateur FROM movements m JOIN articles a ON a.id=m.article_id WHERE m.client=?"
    params = [CLIENT]

    if f_fourn != "Tous":
        sql_m += " AND m.fournisseur=?"
        params.append(f_fourn)
    if f_eq != "Tous":
        sql_m += " AND m.equipe=?"
        params.append(f_eq)

    sql_m += " ORDER BY m.id DESC"
    movs = query(sql_m, tuple(params))
    if movs:
        st.dataframe(pd.DataFrame([dict(m) for m in movs]), use_container_width=True, hide_index=True)

    st.markdown("</div>", unsafe_allow_html=True)

# ---------------------------------------------------------
# ⚙️ RUBRIQUE 5: CONFIGURATION (ADMIN SEULEMENT)
# ---------------------------------------------------------
with tabs[4]:
    if can(ROLE, "config"):
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        st.subheader("Configuration Système (Accès Admin)")

        col_a, col_b, col_c = st.columns(3)

        # 1. Articles & Quantités
        with col_a:
            st.markdown("##### Gérer les Articles")
            with st.form("add_art_form"):
                new_art = st.text_input("Nom de l'Article")
                init_qty = st.number_input("Stock Initial", min_value=0, value=0)
                if st.form_submit_button("Ajouter Article"):
                    if new_art:
                        art_id = execute("INSERT INTO articles (name) VALUES (?)", (new_art,))
                        if init_qty > 0:
                            set_stock(CLIENT, art_id, init_qty)
                            add_movement(CLIENT, art_id, "AJUST_POS", init_qty, "INIT", CURRENT_USER["username"], "Stock Initial")
                        st.success("Article ajouté !")
                        st.rerun()

        # 2. Fournisseurs
        with col_b:
            st.markdown("##### Gérer les Fournisseurs")
            with st.form("add_fourn_form"):
                new_f = st.text_input("Nom du Fournisseur")
                if st.form_submit_button("Ajouter Fournisseur"):
                    if new_f:
                        execute("INSERT INTO fournisseurs (name) VALUES (?)", (new_f,))
                        st.success("Fournisseur ajouté !")
                        st.rerun()

        # 3. Équipes
        with col_c:
            st.markdown("##### Gérer les Équipes")
            with st.form("add_eq_form"):
                new_e = st.text_input("Nom de l'Équipe")
                if st.form_submit_button("Ajouter Équipe"):
                    if new_e:
                        execute("INSERT INTO equipes (name) VALUES (?)", (new_e,))
                        st.success("Équipe ajoutée !")
                        st.rerun()

        st.markdown("---")
        st.subheader("🛠️ Ajustement Manuel du Stock")
        with st.form("manual_adjust_form"):
            adj_art = st.selectbox("Article à ajuster", active_names("articles"))
            adj_qty = st.number_input("Nouvelle Quantité exacte en Stock", min_value=0, value=0)
            if st.form_submit_button("Appliquer l'ajustement"):

                def do_adjust():
                    a_id = article_id_by_name(adj_art)
                    old_qty = current_stock(CLIENT, a_id)
                    diff = adj_qty - old_qty
                    set_stock(CLIENT, a_id, adj_qty)
                    m_type = "AJUST_POS" if diff >= 0 else "AJUST_NEG"
                    add_movement(CLIENT, a_id, m_type, abs(diff), "MANUAL", CURRENT_USER["username"], "Ajustement Manuel Admin")
                    st.success("Stock ajusté !")

                confirm_action_dialog(f"Ajuster manuellement le stock de {adj_art} à {adj_qty} unités ?", do_adjust)

        st.markdown("</div>", unsafe_allow_html=True)
    else:
        st.error("Accès réservé exclusivement aux Administrateurs.")
